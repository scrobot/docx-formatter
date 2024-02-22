import os
import re
from PyQt6.QtWidgets import QApplication, QWidget, QLabel, QPushButton, QLineEdit, QVBoxLayout, QHBoxLayout, QFileDialog

import sys
import asyncio
from docx import Document
from tqdm import tqdm


async def clean_paragraphs(doc):
    def process_para(para):
        # Replace tabs with spaces
        para = re.sub(r'\t', ' ', para)

        # Replace double spaces after a period with single spaces
        para = re.sub(r'\.  ', '. ', para)

        # Remove spaces before commas, periods, dashes, and colons
        para = re.sub(r'\s([.,:;-])', r'\1', para)

        # Remove indents equal to two tabs or more
        para = re.sub(r'\s{2,}', ' ', para)

        return para.strip()

    paragraphs = []
    for para in doc.paragraphs:
        para_text = process_para(para.text)

        # Check if the paragraph is a subchapter separator
        if re.match(r'\* \* \*', para_text) or re.match(r'CHAPTER |PART ', para_text.upper()):
            para_text = f"\n{para_text}\n"
            paragraphs.append(para_text)
            continue

        # If the paragraph text is not empty and the previous paragraph ends with a sentence-ending punctuation mark
        if paragraphs and para_text and re.search(r'[.!?]$', paragraphs[-1]):
            paragraphs.append(para_text)
        else:
            # If this is the first paragraph or the previous paragraph does not end with a sentence-ending punctuation mark, concatenate them
            if paragraphs:
                paragraphs[-1] = f"{paragraphs[-1]} {para_text}".strip()
            else:
                paragraphs.append(para_text)

    new_doc = Document()

    for para in paragraphs:
        new_doc.add_paragraph(para.strip())

    return new_doc


async def process_chunk(doc, output_filename):
    cleaned_doc = await clean_paragraphs(doc)
    cleaned_doc.save(output_filename)


def split_document(doc, chunk_size):
    chunks = []
    current_chunk = []

    for para in doc.paragraphs:
        current_chunk.append(para)

        if len(current_chunk) >= chunk_size:
            chunks.append(current_chunk)
            current_chunk = []

    if current_chunk:
        chunks.append(current_chunk)

    return chunks


def create_document_from_chunks(chunks):
    new_doc = Document()

    for chunk in chunks:
        for para in chunk:
            new_doc.add_paragraph(para.text)

    return new_doc


async def format_docx(input_filename, output_filename):
    document = Document(input_filename)
    chunks = split_document(document, chunk_size=100)
    cleaned_chunks = []

    for i, chunk in enumerate(tqdm(chunks, desc="Processing chunks")):
        temp_doc = create_document_from_chunks([chunk])
        cleaned_chunk_filename = f"temp_cleaned_chunk_{i}.docx"
        await process_chunk(temp_doc, cleaned_chunk_filename)
        cleaned_chunks.append(Document(cleaned_chunk_filename))

    final_doc = Document()

    for cleaned_chunk in cleaned_chunks:
        for para in cleaned_chunk.paragraphs:
            final_doc.add_paragraph(para.text)

    final_doc.save(output_filename)

    # delete temp files
    for i in range(len(cleaned_chunks)):
        os.remove(f"temp_cleaned_chunk_{i}.docx")


class App(QWidget):
    def __init__(self):
        super().__init__()

        self.input_label = QLabel('Input file:')
        self.input_field = QLineEdit()
        self.input_button = QPushButton('Browse')
        self.output_label = QLabel('Output file:')
        self.output_field = QLineEdit()
        self.output_button = QPushButton('Browse')
        self.format_button = QPushButton('Format')

        input_layout = QHBoxLayout()
        input_layout.addWidget(self.input_label)
        input_layout.addWidget(self.input_field)
        input_layout.addWidget(self.input_button)
        output_layout = QHBoxLayout()
        output_layout.addWidget(self.output_label)
        output_layout.addWidget(self.output_field)
        output_layout.addWidget(self.output_button)
        button_layout = QHBoxLayout()
        button_layout.addWidget(self.format_button)
        main_layout = QVBoxLayout()
        main_layout.addLayout(input_layout)
        main_layout.addLayout(output_layout)
        main_layout.addLayout(button_layout)
        self.setLayout(main_layout)

        self.input_button.clicked.connect(self.browse_input)
        self.output_button.clicked.connect(self.browse_output)
        self.format_button.clicked.connect(self.format_doc)

    def browse_input(self):
        file_path, _ = QFileDialog.getOpenFileName(self, 'Open Document', '', 'Word Document (*.docx)')
        if file_path:
            self.input_field.setText(file_path)

    def browse_output(self):    
        file_path, _ = QFileDialog.getSaveFileName(self, 'Save Document', '', 'Word Document (*.docx)')
        if file_path:
            self.output_field.setText(file_path)

    def format_doc(self):
        input_path = self.input_field.text()
        output_path = self.output_field.text()

        loop = asyncio.get_event_loop()
        loop.run_until_complete(format_docx(input_path, output_path))


if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = App()
    window.show()
    sys.exit(app.exec())
